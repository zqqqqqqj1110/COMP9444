from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(r"D:\AirSim\rl_drone_navigation")
WORK = ROOT / "tmp" / "deliverables" / "report"
TEMPLATE = WORK / "template-transitional.docx"
ASSETS = WORK / "assets"
OUTPUT = ROOT / "COMP9444_Project_Report_Autonomous_Drone_Navigation.docx"


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_margins(cell, top=55, start=55, bottom=55, end=55) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "3")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "7A7A7A")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_columns(section, count: int, space_twips: int = 360) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols_node = cols[0]
    else:
        cols_node = OxmlElement("w:cols")
        sect_pr.append(cols_node)
    cols_node.set(qn("w:num"), str(count))
    cols_node.set(qn("w:space"), str(space_twips))


def set_keep(paragraph, keep_next=False, keep_lines=False) -> None:
    fmt = paragraph.paragraph_format
    fmt.keep_with_next = keep_next
    fmt.keep_together = keep_lines
    fmt.widow_control = True


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(9.2)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(1.2)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(10.5)
    normal.paragraph_format.widow_control = True

    title = doc.styles["paper title"]
    title.font.name = "Times New Roman"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title.font.size = Pt(22)
    title.font.bold = False
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(10)
    h1.font.bold = False
    h1.font.italic = False
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(5)
    h1.paragraph_format.space_after = Pt(2)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(9.3)
    h2.font.bold = False
    h2.font.italic = True
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(3)
    h2.paragraph_format.space_after = Pt(1)
    h2.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"] if "Caption" in [style.name for style in doc.styles] else doc.styles.add_style(
        "Caption", WD_STYLE_TYPE.PARAGRAPH
    )
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(7.8)
    caption.font.italic = False
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(1)
    caption.paragraph_format.space_after = Pt(3)
    caption.paragraph_format.keep_with_next = True


def add_text(doc: Document, text: str, *, bold_lead: str | None = None):
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        paragraph.add_run(bold_lead).bold = True
        paragraph.add_run(text[len(bold_lead) :])
    else:
        paragraph.add_run(text)
    set_keep(paragraph)
    return paragraph


def add_equation(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(10)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(8.5)
    run.italic = True
    set_keep(paragraph, keep_lines=True)
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep(paragraph, keep_next=True)
    return paragraph


def add_figure(doc: Document, image: Path, caption: str, width_inches: float = 3.15) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.add_run().add_picture(str(image), width=Inches(width_inches))
    set_keep(paragraph, keep_next=True, keep_lines=True)
    cap = doc.add_paragraph(caption, style="Caption")
    set_keep(cap, keep_next=False, keep_lines=True)


def add_training_table(doc: Document) -> None:
    caption = doc.add_paragraph("TABLE I. CONTROLLED TRAINING CONFIGURATION", style="Caption")
    set_keep(caption, keep_next=True, keep_lines=True)
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [0.68, 1.18, 0.62, 0.70]
    headers = ["Method", "Task exposure", "Final LR", "Selected"]
    values = [
        ["DQN-S", "45k final route", "1e-4", "40.0k"],
        ["PPO-S", "45k final route", "1e-4", "42.5k"],
        ["PPO-C", "5k/10k/30k", "7.5e-5", "20.0k S3"],
    ]
    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        cell.width = Inches(widths[col])
        set_cell_shading(cell, "E6E6E6")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(7.1)
    set_repeat_table_header(table.rows[0])
    for row_index, row_values in enumerate(values, start=1):
        for col, value in enumerate(row_values):
            cell = table.cell(row_index, col)
            cell.width = Inches(widths[col])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if col in (0, 1) else WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(7.1)
    set_table_borders(table)
    after = doc.add_paragraph(
        "Selected is the interaction count of the checkpoint chosen by the two-stage sweep; S3 denotes Curriculum Stage 3."
    )
    after.style = doc.styles["Caption"]
    after.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_results_table(doc: Document) -> None:
    caption = doc.add_paragraph("TABLE II. INDEPENDENT FINAL TEST (DETERMINISTIC, N=50)", style="Caption")
    set_keep(caption, keep_next=True, keep_lines=True)
    table = doc.add_table(rows=4, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [0.92, 0.42, 0.42, 0.42, 0.58, 0.56]
    headers = ["Method", "Succ.", "Coll.", "TO", "Reward", "Dist."]
    values = [
        ["DQN-S", "16%", "72%", "14%", "-18.5", "4.47"],
        ["PPO-S", "98%", "2%", "0%", "141.1", "2.01"],
        ["PPO-C", "68%", "2%", "30%", "71.2", "14.66"],
    ]
    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        cell.width = Inches(widths[col])
        set_cell_shading(cell, "E6E6E6")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(7.2)
    set_repeat_table_header(table.rows[0])
    for row_index, row_values in enumerate(values, start=1):
        for col, value in enumerate(row_values):
            cell = table.cell(row_index, col)
            cell.width = Inches(widths[col])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if col == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(7.2)
            if row_index == 2:
                run.bold = True
    set_table_borders(table)
    after = doc.add_paragraph("Succ./Coll./TO are success, collision, and timeout rates; Dist. is mean final distance in metres.")
    after.style = doc.styles["Caption"]
    after.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_policy_mode_table(doc: Document) -> None:
    caption = doc.add_paragraph("TABLE III. PPO POLICY-MODE DIAGNOSTIC (N=50)", style="Caption")
    set_keep(caption, keep_next=True, keep_lines=True)
    table = doc.add_table(rows=5, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [0.62, 0.60, 0.45, 0.45, 0.45, 0.58]
    headers = ["Method", "Mode", "Succ.", "Coll.", "TO", "Reward"]
    values = [
        ["PPO-S", "Det.", "98%", "2%", "0%", "141.1"],
        ["PPO-S", "Stoch.", "66%", "34%", "0%", "77.7"],
        ["PPO-C", "Det.", "68%", "2%", "30%", "71.2"],
        ["PPO-C", "Stoch.", "46%", "38%", "16%", "41.5"],
    ]
    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        cell.width = Inches(widths[col])
        set_cell_shading(cell, "E6E6E6")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(6.9)
    set_repeat_table_header(table.rows[0])
    for row_index, row_values in enumerate(values, start=1):
        for col, value in enumerate(row_values):
            cell = table.cell(row_index, col)
            cell.width = Inches(widths[col])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if col in (0, 1) else WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(6.9)
            if row_index == 1:
                run.bold = True
    set_table_borders(table)
    after = doc.add_paragraph(
        "Det. selects the highest-probability action; Stoch. samples from the learned categorical distribution."
    )
    after.style = doc.styles["Caption"]
    after.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_references(doc: Document) -> None:
    refs = [
        (
            "[1] S. Shah, D. Dey, C. Lovett, and A. Kapoor, "
            '"AirSim: High-Fidelity Visual and Physical Simulation for Autonomous Vehicles," '
            "in Field and Service Robotics, 2018, pp. 621-635."
        ),
        (
            "[2] V. Mnih et al., "
            '"Human-level control through deep reinforcement learning," '
            "Nature, vol. 518, pp. 529-533, 2015."
        ),
        (
            "[3] J. Schulman, F. Wolski, P. Dhariwal, A. Radford, and O. Klimov, "
            '"Proximal Policy Optimization Algorithms," arXiv:1707.06347, 2017.'
        ),
        (
            "[4] Y. Bengio, J. Louradour, R. Collobert, and J. Weston, "
            '"Curriculum Learning," in Proc. ICML, 2009, pp. 41-48.'
        ),
        (
            "[5] Microsoft Research, "
            '"AirSim," GitHub repository. Available: https://github.com/microsoft/AirSim.'
        ),
        (
            "[6] R. S. Sutton and A. G. Barto, "
            "Reinforcement Learning: An Introduction, 2nd ed. MIT Press, 2018."
        ),
        (
            "[7] H. van Hasselt, A. Guez, and D. Silver, "
            '"Deep Reinforcement Learning with Double Q-learning," in Proc. AAAI, 2016, pp. 2094-2100.'
        ),
    ]
    for text in refs:
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(9.2)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(8)
        set_keep(p, keep_lines=True)


def add_plain_section_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.runs[0]
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def build() -> None:
    doc = Document(TEMPLATE)
    clear_document_body(doc)
    configure_styles(doc)

    first = doc.sections[0]
    first.page_width = Cm(21.0)
    first.page_height = Cm(29.7)
    first.left_margin = Cm(1.58)
    first.right_margin = Cm(1.58)
    first.top_margin = Cm(0.85)
    first.bottom_margin = Cm(1.75)
    first.header_distance = Cm(1.0)
    first.footer_distance = Cm(1.0)
    set_columns(first, 1)

    title = doc.add_paragraph("Autonomous Drone Navigation Using Deep Reinforcement Learning", style="paper title")
    set_keep(title, keep_next=True, keep_lines=True)

    authors = doc.add_paragraph("COMP9444 Project Team")
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.paragraph_format.space_after = Pt(5)
    authors_run = authors.runs[0]
    authors_run.font.name = "Times New Roman"
    authors_run.font.size = Pt(10.5)
    set_keep(authors, keep_next=True, keep_lines=True)

    abstract = doc.add_paragraph()
    abstract.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract.paragraph_format.space_after = Pt(2)
    abstract.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    abstract.paragraph_format.line_spacing = Pt(9.5)
    lead = abstract.add_run("Abstract—")
    lead.bold = True
    lead.italic = True
    body = abstract.add_run(
        "This project develops a visual deep reinforcement learning system for autonomous multirotor navigation in Microsoft AirSim. "
        "The drone observes an 84 x 84 front depth image together with relative goal position and velocity, and chooses among six discrete body-frame actions without an obstacle map. "
        "We implement a vanilla Deep Q-Network (DQN) baseline, a stabilised Proximal Policy Optimisation (PPO) agent trained directly on the final route, and PPO trained through a three-stage distance curriculum. "
        "All methods consume 45,000 environment interactions and are selected by validation before an independent 50-episode test. "
        "Deterministic PPO trained from scratch achieves 98% success with 2% collision, compared with 16%/72% for DQN and 68%/2% for curriculum PPO. "
        "The results show that stable on-policy optimisation can learn a repeatable depth-driven detour, while the curriculum reduces collisions but introduces a 30% timeout rate. "
        "The evidence is limited to one route and one training seed, so cross-route generalisation remains future work."
    )
    for run in (lead, body):
        run.font.name = "Times New Roman"
        run.font.size = Pt(8.7)
    set_keep(abstract, keep_next=True, keep_lines=True)

    keywords = doc.add_paragraph()
    keywords.paragraph_format.space_after = Pt(3)
    keywords.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    keywords.paragraph_format.line_spacing = Pt(9.5)
    label = keywords.add_run("Keywords—")
    label.bold = True
    label.italic = True
    values = keywords.add_run(
        "AirSim, autonomous drone navigation, deep reinforcement learning, DQN, PPO, curriculum learning."
    )
    for run in (label, values):
        run.font.name = "Times New Roman"
        run.font.size = Pt(8.7)
    set_keep(keywords, keep_next=True, keep_lines=True)

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    body_section.page_width = Cm(21.0)
    body_section.page_height = Cm(29.7)
    body_section.left_margin = Cm(1.58)
    body_section.right_margin = Cm(1.58)
    body_section.top_margin = Cm(1.4)
    body_section.bottom_margin = Cm(1.75)
    body_section.header_distance = Cm(1.0)
    body_section.footer_distance = Cm(1.0)
    set_columns(body_section, 2, 360)

    add_heading(doc, "INTRODUCTION")
    add_text(
        doc,
        "Autonomous flight requires a controller to convert high-dimensional visual observations into safe actions under uncertain geometry. "
        "Hand-engineered planners usually depend on maps, calibrated sensing, or explicit obstacle models; these assumptions are difficult to maintain in cluttered and changing environments. "
        "Deep reinforcement learning (DRL) offers an alternative in which perception and action selection are learned jointly from simulator interaction."
    )
    add_text(
        doc,
        "We study point-to-point navigation in the AirSimNH neighbourhood scene. The agent must travel approximately 32.55 m, avoid houses and vegetation, remain within a safe altitude band, and stop within a 2 m goal radius. "
        "The central question is whether PPO learns this visual route more reliably than a vanilla DQN baseline, and whether a progressive target-distance curriculum improves PPO under the same interaction budget."
    )
    add_text(
        doc,
        "This question is separated into four measurable objectives: (1) learn a policy from online visual interaction rather than stored obstacle coordinates; (2) complete the route with few collisions or altitude violations; (3) compare algorithms under an equal 45,000-step interaction budget; and (4) determine whether the learned action distribution is robust enough for deterministic deployment. "
        "Success rate is therefore the primary measure, while collision, timeout, return, path efficiency, final distance, and perceived clearance explain why a policy succeeds or fails."
    )
    add_text(
        doc,
        "Our work is the AirSim-to-Gymnasium environment, visual/state preprocessing, common CNN policy interface, reward and safety design, stable PPO modifications, three-stage transfer procedure, per-step diagnostics, checkpoint selection, and controlled evaluation. "
        "AirSim, DQN, PPO, and the curriculum principle originate from prior work [1]-[4]; the implementation and experiments reported here are our own."
    )

    add_heading(doc, "RELATED WORK")
    add_text(
        doc,
        "AirSim provides an Unreal Engine-based visual and physical simulator for autonomous vehicles, including multirotor dynamics, cameras, collision queries, and programmatic control [1], [5]. "
        "Its repeatability and realistic sensing make it suitable for training without risking physical hardware, although simulator performance alone does not establish real-world transfer."
    )
    add_text(
        doc,
        "DQN demonstrated that convolutional networks can learn value functions directly from image observations in discrete-action tasks [2]. "
        "Its replay buffer and target network improve stability, but off-policy value learning can be sensitive to sparse terminal rewards and correlated failure modes. "
        "PPO instead updates a stochastic actor-critic using a clipped surrogate objective, permitting multiple minibatch epochs while limiting destructive policy changes [3]. "
        "Curriculum learning orders training tasks from easier to harder [4]; here, target distance is increased across stages. "
        "Unlike these foundational benchmarks, our comparison holds the visual encoder, action space, reward, final route, and total interactions constant."
    )
    add_text(
        doc,
        "Prior work supplies the algorithmic principles, not a ready-made solution to this route. DQN was originally demonstrated on stacked game frames, whereas our state combines one metric depth frame with explicit target and velocity features. "
        "PPO is normally presented as a general optimisation procedure rather than an AirSim navigation policy. "
        "Consequently, the main engineering and research decisions in this project are the observation fusion, safety-aware reward, simulator reset handling, curriculum transfer, and evaluation protocol."
    )

    add_heading(doc, "RL TASK AND METHODS")
    add_heading(doc, "Task Formulation and Challenges", level=2)
    add_text(
        doc,
        "The environment is treated as a finite-horizon partially observable Markov decision process. At step t the observation o_t=(I_t, g_t, v_t) contains a depth image I_t, relative goal vector g_t, and linear velocity v_t; the agent selects one of six actions and receives reward r_t. "
        "The transition distribution is produced by AirSim physics and collision geometry. A discount factor of 0.99 values immediate safety while retaining the long-horizon benefit of reaching the goal."
    )
    add_text(
        doc,
        "Because experience is generated online, conventional dataset properties such as sample count, classes, and train/test labels do not apply. The effective data are temporally correlated transitions collected by the current behaviour policy. "
        "The task is challenging because a single forward frame cannot reveal occluded geometry or motion history, progress rewards can favour locally direct motion toward an obstructed goal, and a collision may end an episode before delayed avoidance behaviour receives useful credit. "
        "The NED coordinate convention also means that more negative z values are higher, so altitude constraints and vertical actions require careful sign handling."
    )

    add_heading(doc, "Observation, Actions, and Reward", level=2)
    add_text(
        doc,
        "There is no offline dataset or class label. AirSim generates one transition per interaction. "
        "The visual input is a front DepthPerspective frame clipped to 40 m, converted to obstacle intensity 1-depth/40, resized to 84 x 84, and stored as a single normalised channel. "
        "A six-dimensional state appends scaled relative goal position (dx/50, dy/50, dz/10) and velocity (vx/10, vy/10, vz/10). "
        "The policy receives neither a map nor obstacle coordinates; collision state is used only for reward, termination, and logging."
    )
    add_text(
        doc,
        "Inverse-depth conversion makes nearby surfaces bright and numerically prominent while compressing empty long-range space. Resizing reduces the simulator frame to 7,056 scalar pixels, small enough for repeated CNN updates, and the separate low-dimensional state prevents the network from having to infer the target direction or speed from vision. "
        "The trade-off is information loss: thin branches, distant obstacles, and side obstacles outside the front camera can be missed. The observation is therefore suitable for reactive navigation but is not a complete world model."
    )
    add_text(
        doc,
        "The discrete actions are forward, left, right, up, down, and hover. Each command lasts 0.35 s at 2 m/s horizontally or 1 m/s vertically. "
        "The shaped reward is r = -0.05 + 2(d_prev-d) - 0.25|z-z_target| plus a boundary-margin term. "
        "Reaching the goal adds +100; collision or altitude violation adds -100; a 150-step timeout adds -25. "
        "Episodes terminate on success, a new collision, altitude outside -10 <= z <= -1 m in NED coordinates, or timeout."
    )
    add_equation(
        doc,
        "r_t = -0.05 + 2(d_(t-1)-d_t) - 0.25|z_t-z_g| - p_margin + r_terminal"
    )
    add_text(
        doc,
        "Here p_margin grows linearly when the drone is within 1.5 m of either altitude boundary. The progress term supplies dense feedback, the time cost discourages hovering and unnecessarily long paths, and the terminal terms make safety violations much more expensive than one favourable movement step. "
        "Reward components, position, depth, selected action probability, and terminal cause are logged for every evaluation step, allowing a poor return to be traced to collision, altitude control, timeout, or inefficient motion."
    )

    add_heading(doc, "Shared Visual Encoder and DQN", level=2)
    add_text(
        doc,
        "All agents use three ReLU convolutional layers: 32 filters (8 x 8, stride 4), 64 filters (4 x 4, stride 2), and 64 filters (3 x 3, stride 1). "
        "The resulting 3,136 visual features are concatenated with the six navigation features. "
        "DQN maps the 3,142-vector through a 512-unit ReLU layer to six Q-values. "
        "It uses gamma=0.99, Adam at 1e-4, batch size 32, a 50,000-transition uniform replay buffer, Huber TD loss, gradient clipping, a target update every 1,000 interactions, and epsilon decay from 1.0 to 0.05 over 50,000 steps. "
        "It intentionally omits Double, Dueling, prioritised replay, and n-step extensions, so it is a transparent baseline rather than an optimised DQN."
    )
    add_equation(
        doc,
        "L_DQN = E[Huber(Q_theta(s_t,a_t) - (r_t + gamma max_a Q_target(s_(t+1),a)))]"
    )
    add_text(
        doc,
        "Learning starts after 1,000 transitions. One uniformly sampled minibatch is used per environment interaction, and the target network is copied from the policy network every 1,000 interactions. "
        "At the end of a 45,000-step run, linear exploration is still approximately epsilon=0.145 because the decay horizon is 50,000 steps; evaluation disables exploration and takes argmax Q. "
        "This design gives DQN broad state coverage but also fills much of the replay buffer with early collision-heavy behaviour."
    )

    add_heading(doc, "Stable PPO and Curriculum", level=2)
    add_text(
        doc,
        "PPO feeds the shared representation through a 512-unit layer, per-sample layer normalisation, and Tanh. "
        "Separate linear actor and critic heads output six logits and one state value. "
        "Orthogonal initialisation, reward scaling by 0.1, Huber value loss, advantage normalisation, gradient clipping, and a linear entropy schedule from 0.01 to 0.001 address the activation saturation observed in an earlier implementation. "
        "Training uses gamma=0.99, GAE lambda=0.95, clip coefficient 0.2, 500-step rollouts, batch size 64, four epochs, and Adam at 1e-4."
    )
    add_equation(
        doc,
        "L_clip = E[min(ratio_t A_t, clip(ratio_t, 1-epsilon, 1+epsilon) A_t)]"
    )
    add_text(
        doc,
        "Generalised Advantage Estimation trades variance against bias through lambda=0.95. For each rollout, scaled rewards are bootstrapped with critic values, advantages are normalised, and four shuffled minibatch epochs update the actor and critic. "
        "The clipped probability ratio constrains sudden policy changes, while entropy encourages early action diversity and decays so that late training can form a sharper deployment policy. "
        "Layer normalisation is applied before Tanh so large fused CNN/state activations do not drive all 512 hidden units into saturated regions with near-zero gradients."
    )
    add_text(
        doc,
        "Scratch PPO receives all 45,000 interactions on the final 33 m task. Curriculum PPO uses 5,000 interactions to a 10 m target, 10,000 to a 23 m target, and 30,000 to the final target; maximum episode lengths are 70, 110, and 150. "
        "The selected policy weights transfer between stages while Adam state is reset; Stage 3 uses a 7.5e-5 learning rate. "
        "This equalises total interactions but not final-task exposure, which is central to interpreting the result."
    )
    add_text(
        doc,
        "The curriculum targets are (95.190, -14.491, -3.0), (107.635, -10.842, -3.0), and (117.756, -19.034, -3.0), all from the fixed start (85.413, -15.334, -3.0). "
        "Each stage therefore extends the required displacement while retaining previously encountered geometry. "
        "Transfer is allowed only through model parameters; resetting Adam prevents stale momentum and second-moment estimates from dominating the changed task."
    )

    add_figure(
        doc,
        ASSETS / "training_success.png",
        "Fig. 1. Training success is highly variable. Curriculum stages before 15k interactions use easier targets; only Stage 3 matches the final route.",
    )

    add_heading(doc, "EXPERIMENTAL SETUP")
    add_text(
        doc,
        "Experiments use AirSimNH with start (85.413, -15.334, -3.0) and final target (117.756, -19.034, -3.0). "
        "All three methods use Seed 7, 45,000 new environment interactions, the same observation/action/reward definitions, and a 150-step final-task limit. "
        "Measured training times were 7.41 h for DQN, 6.50 h for Scratch PPO, and 6.51 h for all curriculum stages combined."
    )
    add_training_table(doc)
    add_text(
        doc,
        "To reduce checkpoint-selection bias, every checkpoint is first evaluated for five episodes with the training seed; the top three are then evaluated for 30 episodes with Seed 10007. "
        "The selected model is tested once for 50 fresh episodes with Seed 20007. "
        "This seed changes stochastic simulator and policy sampling but not the route geometry. "
        "The primary deployment metric is deterministic success; secondary metrics are collision, timeout, altitude violation, reward, steps, path length, final distance, and minimum perceived depth. "
        "PPO is additionally evaluated stochastically to diagnose the learned action distribution."
    )
    add_text(
        doc,
        "Checkpoint ranking is lexicographic: higher success first, then lower unsafe rate, lower final distance, and fewer steps. Checkpoints are saved every 2,500 interactions, so selection can recover an earlier policy when continued optimisation degrades behaviour. "
        "The selected files are dqn_step_0040000.pt, ppo_step_0042500.pt, and the Stage 3 ppo_step_0020000.pt. "
        "The final Seed 20007 test is not reused for model choice."
    )
    add_heading(doc, "Evaluation Measures", level=2)
    add_text(
        doc,
        "Success is the fraction of episodes ending within 2 m of the target. Collision and altitude rates quantify unsafe terminal events, while timeout identifies policies that remain safe but fail to finish. "
        "Mean steps and path length measure efficiency; mean final distance distinguishes near misses from severe failures; minimum front depth is a proxy for obstacle clearance. "
        "Return is reported because it reflects the optimised objective, but it is not used alone: shaped reward can increase even when operational success does not."
    )
    add_text(
        doc,
        "For PPO, deterministic evaluation uses argmax over the six logits and approximates deployment with repeatable actions. Stochastic evaluation samples the categorical policy and probes whether non-maximal actions remain safe. "
        "The two modes use the same weights and environment; any gap is therefore evidence about probability mass in the actor, not additional training or a different checkpoint."
    )
    add_heading(doc, "Reproducibility and Integrity Checks", level=2)
    add_text(
        doc,
        "Each run is isolated under experiments/airsimnh/<algorithm>/<run>, with separate models and results directories. The launcher starts the requested AirSim executable, waits for the RPC port, performs a clean-spawn smoke test, runs training or evaluation, records wall-clock time, and closes the scene on completion. "
        "At reset, the environment teleports the multirotor to the configured NED start, zeros motion, and records the collision timestamp. A collision flag inherited from scene spawn geometry is ignored only as a stale baseline; any newer timestamp terminates the episode. This avoids rewarding an agent for an invalid initial state while preserving subsequent collision detection."
    )
    add_text(
        doc,
        "Training logs store episode, cumulative interaction count, return, length, success, collision, altitude violation, final distance, exploration statistics, and PPO optimisation diagnostics. Evaluation adds complete per-step trajectories containing before/after position, depth clearance, action name and probability, every reward component, cumulative return, and terminal cause. "
        "Model files contain network parameters and step metadata, while the selected checkpoint is copied to a clearly named deterministic-best file. These outputs allow the tables and plots in this report to be regenerated from CSV records rather than transcribed manually."
    )

    add_heading(doc, "RESULTS")
    add_heading(doc, "Training Behaviour", level=2)
    add_text(
        doc,
        "Episode outcomes during training are noisy because DQN follows epsilon-greedy exploration and PPO samples actions with a non-zero entropy bonus. Figure 1 should therefore be read as learning dynamics rather than final deployment performance. "
        "DQN does not establish a sustained high-success region. Scratch PPO develops successful final-route behaviour late in training, while curriculum success rises on the shorter stages and changes again when the target moves. "
        "This target-dependent discontinuity shows why Stage 1 and Stage 2 episodes cannot be counted as final-route successes."
    )
    add_results_table(doc)
    add_figure(
        doc,
        ASSETS / "final_outcomes.png",
        "Fig. 2. Independent deterministic outcomes. Rates are reported separately because a terminal event at the step limit can overlap timeout bookkeeping.",
    )
    add_text(
        doc,
        "Scratch PPO succeeds in 49/50 episodes (98%; Wilson 95% CI 89.5-99.6), versus 8/50 for DQN (16%; 8.3-28.5) and 34/50 for curriculum PPO (68%; 54.2-79.2). "
        "It also obtains the highest reward (141.1), shortest mean completion time (52.28 steps), and a 1.92 m mean minimum depth. "
        "Its 38.27 m path is 17.6% longer than the straight-line distance, consistent with a deliberate lateral detour rather than direct pursuit."
    )
    add_text(
        doc,
        "DQN collides in 72% of episodes, reaches a minimum perceived depth of only 0.37 m on average, and often alternates direct motion with hover near the obstacle. "
        "Curriculum PPO has the same low 2% deterministic collision rate as Scratch PPO, but times out in 30% of episodes and therefore ends 14.66 m from the target on average. "
        "Its 55.17 m average path indicates inefficient wandering in failed episodes."
    )
    add_figure(
        doc,
        ASSETS / "representative_trajectories.png",
        "Fig. 3. Representative deterministic Episode 1 trajectories. Both PPO policies move laterally before approaching the target; DQN follows a lower-clearance route and collides.",
    )
    add_text(
        doc,
        "The deployment mode matters. Scratch PPO falls from 98% deterministic success to 66% when actions are sampled; curriculum PPO falls from 68% to 46% and its collision rate rises from 2% to 38%. "
        "Thus the argmax policy is repeatable, but probability mass remains on unsafe alternatives. "
        "The reported mean final distance of 2.01 m for Scratch PPO is compatible with a 2 m success radius because it averages 49 near-boundary successes with one collision failure."
    )
    add_policy_mode_table(doc)
    add_heading(doc, "Failure Analysis", level=2)
    add_text(
        doc,
        "Per-step trajectories clarify the aggregate rates. DQN's average 0.37 m minimum depth and 72% collision rate indicate that its value estimates do not consistently trigger lateral avoidance before clearance becomes critical. "
        "Scratch PPO produces a repeatable side movement before returning toward the target, explaining both its 38.27 m path and high success. Its only deterministic failure keeps the average final distance slightly above the success radius despite 49 successful episodes."
    )
    add_figure(
        doc,
        ASSETS / "action_distribution.png",
        "Fig. 4. Aggregate deterministic action distribution over all 50 final-test episodes. Percentages are computed from complete per-step trajectory logs.",
    )
    add_text(
        doc,
        "Action usage supports the trajectory interpretation. Scratch PPO allocates 75.6% of its 2,614 actions to forward motion, 16.5% to left, and 7.8% to right, with no vertical or hover actions; it has learned a compact horizontal detour at the fixed target altitude. "
        "Curriculum PPO uses only 54.5% forward and 30.4% left, plus 8.0% vertical actions and 0.4% hover, consistent with its longer paths and timeouts. "
        "DQN is the least decisive: only 41.1% forward, with 14.2% hover and 25.6% combined vertical actions. These descriptive counts do not prove causality, but they identify the behavioural differences behind the endpoint metrics."
    )
    add_text(
        doc,
        "Curriculum PPO is safer than DQN but less decisive than Scratch PPO. Deterministic failures are predominantly timeouts rather than collisions, and the 14.66 m final distance shows that they are not merely one-step near misses. "
        "When sampled, both PPO policies expose substantially more collisions, demonstrating that evaluation of only the most likely action can conceal risky alternatives. "
        "For an autonomous system, deterministic performance is the correct deployment headline, but stochastic diagnostics remain important evidence about robustness."
    )
    add_heading(doc, "Efficiency and Computational Cost", level=2)
    add_text(
        doc,
        "At 2 m/s for 0.35 s, an ideal forward command covers about 0.7 m. The 32.55 m straight-line displacement therefore has a nominal lower bound near 47 commands before accounting for acceleration, sensing, and avoidance. "
        "Scratch PPO averages 52.28 steps on the test, only about six commands above this idealised bound, while following a 38.27 m path. DQN and Curriculum PPO average 86.80 and 83.46 steps respectively; because their averages include collisions and timeouts, path length and step count must be interpreted jointly with success."
    )
    add_text(
        doc,
        "The shared encoder has 71,840 trainable convolutional parameters. Including fusion and output layers, DQN has approximately 1.684 million parameters and PPO 1.686 million, the small difference coming from LayerNorm and separate actor/critic heads. "
        "Thus PPO's advantage is not explained by a materially larger network. Measured throughput is about 1.69 interactions/s for DQN and 1.92 interactions/s for both PPO regimes, including AirSim control and image transfer. Wall-clock differences are modest relative to the accuracy gap."
    )
    add_text(
        doc,
        "Checkpoint selection also improves computational efficiency at deployment: the best policies occur at 40.0k, 42.5k, and 20.0k Stage 3 interactions rather than necessarily at the final saved model. "
        "Continuing training after a useful policy can alter the actor or Q-function, so retaining intermediate checkpoints is essential in non-monotonic RL. The sweep does not create new capability; it identifies which already-trained policy should be evaluated and deployed."
    )

    add_heading(doc, "DISCUSSION AND CONCLUSION")
    add_text(
        doc,
        "The controlled comparison supports three conclusions. First, PPO is substantially better matched to this visual navigation task than the implemented vanilla DQN: clipped on-policy updates, GAE, and actor-critic learning produce a stable detour instead of repeatedly bootstrapping values from collision-heavy experience. "
        "Second, curriculum learning is not automatically beneficial. Spending one third of the budget on shorter goals reduces final-route exposure, and changing targets creates a distribution shift; the transferred policy is safe when successful but preserves short-horizon behaviour that manifests as timeouts. "
        "Third, deterministic and stochastic evaluation answer different questions: argmax measures deployment performance, while sampling reveals residual uncertainty."
    )
    add_text(
        doc,
        "The comparison also illustrates the difference between equal interaction cost and equal task exposure. All methods consume 45,000 transitions, but Scratch PPO receives 50% more final-route experience than Curriculum PPO (45,000 versus 30,000). "
        "The curriculum result is therefore not evidence that curricula are generally harmful; it shows that this particular distance schedule and fixed budget do not compensate for reduced Stage 3 rehearsal. "
        "A stronger curriculum would interleave earlier and final targets or allocate additional final-route steps while reporting both total and final-task budgets."
    )
    add_text(
        doc,
        "DQN remains a useful baseline despite its low score because it is a recognised image-based discrete-control method [2] and shares the same encoder and action interface. Its result establishes that the environment is not solved by arbitrary CNN value learning under this budget. "
        "Nevertheless, the comparison should not be interpreted as PPO outperforming every DQN variant: Double Q-learning [7], duelling heads, prioritised replay, longer exploration, or frame stacking could improve value estimation. The correct claim is that the implemented stable PPO configuration outperforms the documented vanilla DQN baseline on this controlled route."
    )
    add_heading(doc, "What the Policy Has Learned", level=2)
    add_text(
        doc,
        "The simulator does not provide obstacle coordinates to the agent and the implementation does not build a map. Avoidance decisions are produced from the current inverse-depth pattern, relative target vector, velocity, and learned network weights. "
        "In this sense, the policy learns a visual action rule rather than recording a list of tree or building positions. Collision information affects gradients through reward and termination during training but is unavailable as an advance warning at inference."
    )
    add_text(
        doc,
        "However, a fixed start and goal make relative goal position strongly correlated with location along the route. A sufficiently expressive network can therefore associate particular goal vectors with left or right actions, effectively memorising a route phase even while receiving depth. "
        "The repeatable Scratch PPO detour and concentrated three-action distribution are compatible with both reactive avoidance and route memorisation. Only counterfactual tests, such as moving an obstacle or perturbing the start while keeping similar depth patterns, can separate these explanations. This is the main reason the present result is described as fixed-route navigation."
    )
    add_heading(doc, "Reward Design and Safety", level=2)
    add_text(
        doc,
        "No direct reward is given for selecting a lateral action or maintaining a particular depth clearance. The agent must discover avoidance because a collision produces -100 and removes the future +100 goal opportunity. "
        "Dense progress feedback makes learning feasible, but it can also penalise the first part of a detour when distance temporarily increases. The successful Scratch PPO trajectory shows that its long-term advantage estimates overcome this local conflict; DQN's bootstrap targets do so less consistently."
    )
    add_text(
        doc,
        "The 2 m success radius is deliberately larger than one nominal 0.7 m horizontal command, reducing oscillation around an exact coordinate. Altitude penalties are symmetric in terminal cost but NED signs remain asymmetric in implementation. "
        "The absence of altitude violations in all final tests suggests that the hold and margin terms work on this route, not that vertical safety is solved generally. A deployment-oriented extension should add side sensing, emergency braking, and a rule-based safety layer rather than relying only on learned penalties."
    )
    add_text(
        doc,
        "Strengths of the system include map-free visual control, a fair interaction budget, identical final-task evaluation, explicit safety metrics, checkpoint selection separated from testing, and action/reward trajectory logs that make failure analysis reproducible. "
        "The main weakness is external validity. There is only one training seed, one fixed route, one scene, one front depth frame, and no recurrent memory, yaw, backward action, domain randomisation, or real-drone validation. "
        "The test seed is independent of checkpoint selection but does not create unseen geometry, so these results must not be described as generalisation."
    )
    add_text(
        doc,
        "Internal validity is stronger than external validity. The methods share task definitions and final evaluation, and the 50-episode test gives useful uncertainty intervals: Scratch PPO's 98% success has a Wilson 95% interval of 89.5-99.6%, compared with 54.2-79.2% for Curriculum PPO and 8.3-28.5% for DQN. "
        "However, these intervals measure episode variability for the selected trained policies; they do not account for variation across training seeds. One unusually good or poor Seed 7 optimisation run could therefore change the algorithm ranking."
    )
    add_text(
        doc,
        "Future work should repeat training across at least three to five seeds, perturb safe start and target positions, evaluate a second AirSim environment, and train on multiple routes with weather and texture randomisation. "
        "Algorithmic extensions should compare Double/Dueling DQN, recurrent visual memory, continuous-control PPO or SAC, and a curriculum that preserves final-task rehearsal. "
        "Within the current scope, deterministic Scratch PPO is the recommended policy: it achieves 98% success and 2% collision on the held-out 50-episode fixed-route test."
    )
    add_text(
        doc,
        "A practical next experiment is a route-generalisation matrix: train each method on the current route using three seeds, then evaluate deterministic policies on several safe start/goal perturbations and one unseen scene without further learning. "
        "Reporting mean and standard deviation across seeds, success versus perturbation distance, and collision objects would separate memorisation of a fixed detour from transferable visual obstacle avoidance. "
        "Until that experiment is performed, the defensible conclusion is high fixed-route deployment performance for Scratch PPO, not general autonomous navigation."
    )

    add_plain_section_heading(doc, "REFERENCES")
    add_references(doc)

    # Footer is intentionally minimal and source-derived.
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = True
        paragraph = footer.paragraphs[0]
        paragraph.text = "COMP9444 - Neural Networks and Deep Learning"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(7.5)
            run.font.color.rgb = None

    doc.core_properties.title = "Autonomous Drone Navigation Using Deep Reinforcement Learning"
    doc.core_properties.subject = "COMP9444 Group Project Summary Report"
    doc.core_properties.author = "COMP9444 Project Team"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
